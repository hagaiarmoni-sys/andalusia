"""
Document Generator for Andalusia Travel App
Generates beautiful Word documents with travel itineraries
FINAL VERSION: Restores full feature set including Detailed Driving Segments and Summary Box.
"""

import os
import io
import math
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta

# Import monetization
try:
    from monetization_manager import get_hotel_affiliate_link
except ImportError:
    # Safe fallback if monetization_manager.py is not created yet
    from urllib.parse import quote_plus
    def get_hotel_affiliate_link(h, c, i=None, o=None): 
        return f"https://www.booking.com/searchresults.html?ss={quote_plus(h+' '+c)}"

# DYNAMIC PATHS
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
PHOTOS_DIR = os.path.join(DATA_DIR, 'photos')

# --- ESTIMATION LOGIC ---
def estimate_fuel_cost(total_km):
    """
    Estimates fuel cost in Euros based on distance.
    Assumes: 8 L/100km efficiency and 1.60 EUR/L fuel price.
    Returns: a tuple of (min_cost, max_cost) for a range.
    """
    # Calculation: KM / 100 * Efficiency_L_per_100km * Price_EUR_per_L
    base_cost = (total_km / 100) * 8 * 1.60
    
    # Return a reasonable range based on driving conditions
    min_cost = int(base_cost * 0.85)
    max_cost = int(base_cost * 1.2)
    
    return min_cost, max_cost

def add_hyperlink(paragraph, url, text):
    """Add a working hyperlink to a Word document paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2980B9')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_route_at_a_glance(doc, itinerary, days, ordered_cities, maps_link):
    """
    Adds the 'Route At A Glance' summary box, matching the original good document.
    """
    
    # Calculate Total Driving Stats
    total_km = sum(d.get('driving_km', 0) for d in itinerary)
    total_hours = sum(d.get('driving_hours', 0) for d in itinerary)
    
    avg_km_per_day = round(total_km / days) if days > 0 else 0
    
    min_fuel, max_fuel = estimate_fuel_cost(total_km)
    
    doc.add_heading("🗺️ YOUR ROUTE AT A GLANCE", level=1)
    
    # 1. City list
    cities_str = " → ".join(ordered_cities)
    p = doc.add_paragraph()
    p.add_run("🎯 ").bold = True
    p.add_run(cities_str)

    # 2. Stats
    stats_p = doc.add_paragraph()
    stats_p.paragraph_format.space_before = Pt(6)
    
    stats_p.add_run(f"🚗  Total Driving: {round(total_km)}km (~{total_hours:.1f} hours)    ").bold = True
    stats_p.add_run(f"📊  Average per Day: {avg_km_per_day}km    ").bold = True
    stats_p.add_run(f"⛽  Estimated Fuel Cost: €{min_fuel}-{max_fuel}").bold = True
    
    # 3. Maps Link
    doc.add_paragraph()
    p_map = doc.add_paragraph()
    p_map.add_run("🌍  ").bold = True
    add_hyperlink(p_map, maps_link, "OPEN ROUTE IN GOOGLE MAPS")


def add_driving_segments(doc, itinerary):
    """
    Adds the detailed segment-by-segment driving breakdown table.
    """
    
    # Only show if there is driving in the itinerary
    if not any(d.get('driving_km', 0) > 0 for d in itinerary):
        return

    doc.add_paragraph()
    doc.add_heading("🚗 DRIVING SEGMENTS", level=2)
    doc.add_paragraph("Your scenic drives through Andalusia - times are approximate, not including stops for photos, coffee, or impromptu adventures! ☕📸")
    doc.add_paragraph()
    
    # Create a table for a clean two-column layout (Source -> Destination, Details)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    
    # Add a dummy row to hold the structure
    table.cell(0, 0).text = "Source → Destination"
    table.cell(0, 1).text = "Driving Details"
    
    # Apply a distinct color to the header
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(41, 128, 185) # Blue

    # Add content rows
    for day in itinerary:
        # Check if this day involves a drive to the overnight city
        if day.get("driving_km", 0) > 0 and day.get('overnight_city'):
            source_city = day.get('city')
            dest_city = day.get('overnight_city')
            driving_km = day['driving_km']
            driving_hours = day.get('driving_hours', 0.0)
            
            # Recalculate segment fuel cost estimate
            min_fuel, max_fuel = estimate_fuel_cost(driving_km)
            avg_fuel = int((min_fuel + max_fuel) / 2) # Use average for segment

            # Add a row
            row_cells = table.add_row().cells
            
            # Column 1: Source -> Destination
            p1 = row_cells[0].paragraphs[0]
            p1.add_run(f"{source_city} ━━━━━━━━➤ {dest_city}").bold = True
            
            # Column 2: Details
            p2 = row_cells[1].paragraphs[0]
            p2.add_run(f"📍 {driving_km}km  •  ").bold = True
            p2.add_run(f"⏱️ ~{driving_hours:.1f}h drive  •  ").bold = True
            p2.add_run(f"⛽ ~€{avg_fuel} fuel").bold = True

def add_events_to_doc(doc, result):
    """
    Adds the Events section to the Word document, restoring the original logic.
    (This function remains robust as implemented in the last step)
    """
    # Assuming trip_planner_page.py passes fetched events in 'all_events' key
    all_events = result.get('all_events', [])
    
    # Get a list of unique cities in the itinerary for filtering
    itinerary_cities = set()
    for day in result.get('itinerary', []):
        city = day.get('city')
        if city: itinerary_cities.add(city)
    
    # Filter and sort events to be relevant to the cities visited
    relevant_events = [e for e in all_events if e.get('city') in itinerary_cities]
    relevant_events.sort(key=lambda x: x.get('date', '9999-01-01'))
    
    if not relevant_events:
        return # Skip section if no relevant events are found

    doc.add_page_break()
    doc.add_heading("🎉 Events During Your Trip", level=1)
    doc.add_paragraph("Don't miss these local events happening in the cities on your route!")
    doc.add_paragraph()
    
    for event in relevant_events[:10]: # Limit to 10 for document clarity
        event_name = event.get('name', 'Unknown Event')
        event_date = event.get('date', 'TBD')
        event_location = event.get('location', event.get('city', ''))
        event_type = event.get('type', 'General')
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        
        # Event Header (Green)
        run = p.add_run(f"📅 {event_date}: {event_name}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(46, 204, 113) # Emerald Green

        # Details
        details_p = doc.add_paragraph()
        details_p.paragraph_format.left_indent = Inches(0.2)
        details_p.paragraph_format.space_before = Pt(0)
        details_p.paragraph_format.space_after = Pt(6)
        
        details_p.add_run(f"📍 {event_location}  |  🏷️ {event_type}").italic = True
        
        # Tip for major festivals (Restored from your original code)
        if event_type in ['Festival', 'Carnival', 'Feria'] or 'Feria' in event_name or 'Festival' in event_name:
            tip_para = doc.add_paragraph()
            tip_para.paragraph_format.left_indent = Inches(0.2)
            tip_para.paragraph_format.space_before = Pt(0)
            tip_para.paragraph_format.space_after = Pt(6)
            tip_run = tip_para.add_run('💡 Tip: Book accommodations early - prices increase during festivals!')
            tip_run.font.size = Pt(9)
            tip_run.italic = True
            tip_run.font.color.rgb = RGBColor(241, 196, 15)  # Golden yellow

# --- CORRECTED FUNCTION SIGNATURE ---
def add_final_summary(doc, result, days, prefs, ordered_cities, maps_link):
    """
    Adds the Final Summary table and Google Maps Link section.
    (This function remains robust as implemented in the last step)
    """
    doc.add_page_break()
    doc.add_heading("📊 Trip Summary & Logistics", level=1)
    
    # 1. Total Driving Distance
    total_km = result.get('total_km', 0)
    
    # 2. Preferences (Pace/Budget)
    pace_val = prefs.get('pace', '?').title()
    budget_val = prefs.get('budget', '?').title()
    
    # Summary Table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Medium List 1 Accent 1'
    
    # Populate rows
    rows = [
        ("Trip Type", result.get('trip_type', 'Road Trip')),
        ("Duration", f"{days} Days"),
        ("Cities Visited", ", ".join(ordered_cities)),
        ("Total Driving Distance (Approx)", f"{round(total_km)} km"),
        ("Pace/Budget", f"{pace_val} / {budget_val}"),
        ("Cities Count", f"{len(ordered_cities)} Cities")
    ]
    
    for i, (key, value) in enumerate(rows):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = str(value)

    # Google Maps Link
    doc.add_heading("🗺️ Full Route Link", level=2)
    p = doc.add_paragraph("Click here for the full route on Google Maps: ")
    
    # Use the maps_link passed directly to build_word_doc
    link_url = result.get('maps_link', maps_link)
    add_hyperlink(p, link_url, "Google Maps Route")
    
    doc.add_paragraph()

# --- MAIN GENERATOR FUNCTION ---
def build_word_doc(itinerary, hop_kms, maps_link, ordered_cities, days, prefs, parsed_requests, is_car_mode=False, result=None):
    """Build BEAUTIFUL travel magazine-style Word document (The main function)"""
    
    doc = Document()
    
    # --- COVER PAGE (ENHANCED) ---
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_heading('', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('✈️ YOUR ANDALUSIA\nROAD TRIP ADVENTURE ✈️')
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(41, 128, 185)
    title_run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    start = ordered_cities[0] if ordered_cities else "Start"
    end = ordered_cities[-1] if ordered_cities and len(ordered_cities) > 1 else start
    
    # Route line (similar to old document)
    route_run = subtitle.add_run(f'🚗 {start} → {end} 🚗')
    route_run.font.size = Pt(20)
    route_run.font.color.rgb = RGBColor(231, 76, 60)
    route_run.bold = True
    
    # New detailed information line
    details_p = doc.add_paragraph()
    details_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    start_date_str = result.get('start_date', '??-??-????')
    end_date_str = result.get('end_date', '??-??-????')

    details_p.add_run(f"📅 {days} Days  •  ")
    details_p.add_run(f"🏨 {len(ordered_cities)} Cities  •  ")
    details_p.add_run(f"{prefs.get('budget', '?').title()} Budget  •  ")
    details_p.add_run(f"{start_date_str} → {end_date_str}")
    details_p.add_run('\n')

    # Quote
    quote_p = doc.add_paragraph()
    quote_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_p.paragraph_format.space_before = Pt(12)
    quote_p.add_run("“The world is a book, and those who do not travel read only one page.”").italic = True
    quote_p.add_run(" – Saint Augustine")

    doc.add_page_break()
    
    # --- ROUTE AT A GLANCE (NEW SECTION) ---
    add_route_at_a_glance(doc, itinerary, days, ordered_cities, maps_link)
    
    # --- DETAILED DRIVING SEGMENTS (NEW SECTION) ---
    add_driving_segments(doc, itinerary)
    
    doc.add_page_break() # Start itinerary on a clean page
    
    # --- DAILY ITINERARY ---
    for idx, day in enumerate(itinerary):
        day_num = day.get("day", 0)
        city = day.get("city", "?")
        date_obj = day.get('date_obj')
        
        # Format the date nicely from the date object
        date_str = date_obj.strftime('%A, %B %d') if date_obj else f"Day {day_num}"
        
        # Header
        day_header = doc.add_heading(f"📅 {date_str}: {city}", level=1)
        day_header.runs[0].font.color.rgb = RGBColor(155, 89, 182)
        
        # Driving Info
        if day.get("driving_km", 0) > 0:
            p = doc.add_paragraph()
            p.add_run(f"🚗 Drive: {day['driving_km']}km (~{day.get('driving_hours',0)}h) to {day.get('overnight_city', city)}").bold = True
        
        # Attractions
        doc.add_heading("🎯 Highlights", level=2)
        for city_stop in day.get("cities", []):
            for attr in city_stop.get("attractions", []):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.add_run(f"• {attr.get('name')}").bold = True
                
                # Photo Logic (using dynamic relative paths)
                photo_path = None
                if attr.get('local_photo_path'):
                    # Fix path separators and re-assemble path
                    rel_path = attr['local_photo_path'].replace('\\', '/').split('photos/')[-1]
                    cand = os.path.join(PHOTOS_DIR, rel_path)
                    if os.path.exists(cand): photo_path = cand
                
                if photo_path:
                    try:
                        pic_p = doc.add_paragraph()
                        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = pic_p.add_run()
                        run.add_picture(photo_path, width=Inches(4.5))
                    except Exception as e:
                        pass
                
                if attr.get('description'):
                    doc.add_paragraph(attr['description']).italic = True

        # Hotels with Monetization
        hotels = day.get("hotels", [])
        if hotels:
            doc.add_heading("🏨 Where to Stay", level=2)
            checkin = day.get('date_obj')
            checkout = checkin + timedelta(days=1) if checkin else None
            
            for hotel in hotels[:3]:
                h_p = doc.add_paragraph()
                h_p.add_run(f"{hotel.get('name')}").bold = True
                if hotel.get('rating'): h_p.add_run(f" (⭐ {hotel['rating']})")
                
                link = get_hotel_affiliate_link(hotel.get('name'), day.get('overnight_city', city), checkin, checkout)
                h_p.add_run("  ")
                add_hyperlink(h_p, link, "👉 Check Rates on Booking.com")

        # Restaurants
        lunch = day.get("lunch_restaurant")
        dinner = day.get("dinner_restaurant")
        if lunch or dinner:
             doc.add_heading("🍽️ Dining", level=2)
             if lunch: doc.add_paragraph(f"🥗 Lunch: {lunch.get('name')}")
             if dinner: doc.add_paragraph(f"🍷 Dinner: {dinner.get('name')}")
             
        # Add page break only if not the last day
        if idx < len(itinerary) - 1:
            doc.add_page_break()

    # --- ADD POST-ITINERARY SECTIONS ---
    
    # 1. Events Section
    events_result = result if result is not None else {}
    add_events_to_doc(doc, events_result)

    # 2. Final Summary (Uses direct parameters for reliability)
    add_final_summary(doc, result or {}, days, prefs, ordered_cities, maps_link)

    # Closing
    doc.add_paragraph()
    end = doc.add_paragraph("Generated with ❤️ by Your Andalusia Planner")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio