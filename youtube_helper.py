"""
YouTube Video Helper for Andalusia Travel App
Provides functions to find and add relevant YouTube videos to documents

Usage:
    from youtube_helper import get_video_for_city, add_youtube_section_to_doc
"""

import json
import os
import unicodedata
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# PATH CONFIGURATION
# ============================================================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
YOUTUBE_DB_PATH = os.path.join(DATA_DIR, 'youtube_videos_db.json')

# Cache for loaded videos
_youtube_cache = None


def _load_youtube_db():
    """Load YouTube videos database with caching"""
    global _youtube_cache
    
    if _youtube_cache is not None:
        return _youtube_cache
    
    try:
        if os.path.exists(YOUTUBE_DB_PATH):
            with open(YOUTUBE_DB_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
                _youtube_cache = data.get('destinations', {})
                print(f"✅ Loaded YouTube videos DB: {len(_youtube_cache)} destinations")
                return _youtube_cache
        else:
            print(f"⚠️ YouTube DB not found at: {YOUTUBE_DB_PATH}")
            return {}
    except Exception as e:
        print(f"❌ Error loading YouTube DB: {e}")
        return {}


def normalize_name(name):
    """Normalize city/attraction name for matching"""
    if not name:
        return ""
    name = str(name)
    # Remove accents
    nfd = unicodedata.normalize('NFD', name)
    without_accents = ''.join(char for char in nfd if unicodedata.category(char) != 'Mn')
    return without_accents.lower().strip()


def get_video_for_city(city_name, max_videos=1):
    """
    Get YouTube video(s) for a city
    
    Args:
        city_name: Name of the city (e.g., "Seville", "Granada")
        max_videos: Maximum number of videos to return (default 1)
    
    Returns:
        List of video dicts with keys: title, video_id, watch_url, thumbnail_url, channel
        Returns empty list if no videos found
    """
    youtube_db = _load_youtube_db()
    
    if not youtube_db:
        return []
    
    city_norm = normalize_name(city_name)
    
    # Try exact match first
    for db_city, videos in youtube_db.items():
        if normalize_name(db_city) == city_norm:
            # Filter out sports/irrelevant videos
            filtered = _filter_travel_videos(videos)
            return filtered[:max_videos]
    
    # Try partial match
    for db_city, videos in youtube_db.items():
        if city_norm in normalize_name(db_city) or normalize_name(db_city) in city_norm:
            filtered = _filter_travel_videos(videos)
            return filtered[:max_videos]
    
    return []


def get_video_for_attraction(attraction_name, city_name=None, max_videos=1):
    """
    Get YouTube video for a specific attraction
    
    Args:
        attraction_name: Name of the attraction (e.g., "Alhambra", "Mezquita")
        city_name: Optional city name for better matching
        max_videos: Maximum number of videos to return
    
    Returns:
        List of video dicts
    """
    youtube_db = _load_youtube_db()
    
    if not youtube_db:
        return []
    
    attr_norm = normalize_name(attraction_name)
    
    # Special attraction mappings
    attraction_mappings = {
        'alhambra': ['Alhambra', 'Granada'],
        'mezquita': ['Mezquita', 'Córdoba'],
        'alcazar': ['Alcázar Seville', 'Seville'],
        'caminito del rey': ['Caminito del Rey'],
        'plaza de espana': ['Seville'],
        'giralda': ['Seville'],
        'sierra nevada': ['Sierra Nevada'],
        'el torcal': ['El Torcal'],
        'donana': ['Doñana'],
    }
    
    # Check attraction mappings
    for key, db_keys in attraction_mappings.items():
        if key in attr_norm:
            for db_key in db_keys:
                if db_key in youtube_db:
                    filtered = _filter_travel_videos(youtube_db[db_key])
                    if filtered:
                        return filtered[:max_videos]
    
    # Fallback to city video
    if city_name:
        return get_video_for_city(city_name, max_videos)
    
    return []


def _filter_travel_videos(videos):
    """
    Filter out non-travel videos (sports, news, etc.)
    
    Args:
        videos: List of video dicts
    
    Returns:
        Filtered list with only travel-related videos
    """
    if not videos:
        return []
    
    # Keywords that indicate non-travel content
    exclude_keywords = [
        'laliga', 'fc barcelona', 'real madrid', 'sevilla fc', 'athletic',
        'highlights', 'resumen', 'gol', 'match', 'partido', 'futbol', 'football',
        'basketball', 'padel', 'deportes', 'sports',
        'noticias', 'news', 'rtve', 'polemic',
    ]
    
    # Keywords that indicate travel content (boost these)
    include_keywords = [
        'travel', 'tour', 'drone', '4k', 'walking', 'guide', 'visit',
        'spain', 'andalusia', 'andalucia', 'españa',
        'city', 'village', 'pueblo', 'castle', 'palace', 'beach',
        'rick steves', 'lonely planet',
    ]
    
    filtered = []
    for video in videos:
        title_lower = video.get('title', '').lower()
        channel_lower = video.get('channel', '').lower()
        
        # Check for exclusions
        is_excluded = any(kw in title_lower or kw in channel_lower for kw in exclude_keywords)
        
        if not is_excluded:
            # Prioritize videos with travel keywords
            has_travel_keyword = any(kw in title_lower for kw in include_keywords)
            if has_travel_keyword:
                filtered.insert(0, video)  # Add to front
            else:
                filtered.append(video)
    
    return filtered


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph
    (Duplicated here for standalone use, but should use document_generator's version)
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    part = paragraph.part
    r_id = part.relate_to(
        url, 
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 
        is_external=True
    )
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'CC0000')  # YouTube red
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def add_youtube_section_to_doc(doc, city_name, add_hyperlink_func=None):
    """
    Add a YouTube video section for a city to the document
    
    Args:
        doc: python-docx Document object
        city_name: Name of the city
        add_hyperlink_func: Optional function to add hyperlinks (from document_generator)
    
    Returns:
        True if video was added, False otherwise
    """
    videos = get_video_for_city(city_name, max_videos=1)
    
    if not videos:
        return False
    
    video = videos[0]
    
    # Use provided hyperlink function or our local one
    hyperlink_fn = add_hyperlink_func or add_hyperlink
    
    # Create video section
    video_para = doc.add_paragraph()
    video_para.paragraph_format.left_indent = Inches(0.3)
    video_para.paragraph_format.space_before = Pt(12)
    video_para.paragraph_format.space_after = Pt(6)
    
    # YouTube icon and label
    yt_icon = video_para.add_run('🎬  ')
    yt_icon.font.size = Pt(12)
    
    label = video_para.add_run('Watch a preview:  ')
    label.font.size = Pt(18)
    label.font.color.rgb = RGBColor(52, 73, 94)
    
    # Add clickable link
    watch_url = video.get('watch_url', '')
    video_title = video.get('title', 'Video')
    
    # Truncate long titles
    if len(video_title) > 50:
        video_title = video_title[:47] + '...'
    
    hyperlink_fn(video_para, watch_url, f'▶ {video_title}')
    
    # Channel info
    channel = video.get('channel', '')
    if channel:
        channel_para = doc.add_paragraph()
        channel_para.paragraph_format.left_indent = Inches(0.5)
        channel_run = channel_para.add_run(f'📺 Channel: {channel}')
        channel_run.font.size = Pt(8)
        channel_run.italic = True
        channel_run.font.color.rgb = RGBColor(127, 140, 141)
    
    return True


def add_youtube_for_attraction(doc, attraction_name, city_name=None, add_hyperlink_func=None):
    """
    Add a YouTube video for a specific attraction
    
    Args:
        doc: python-docx Document object
        attraction_name: Name of the attraction
        city_name: Optional city name for fallback
        add_hyperlink_func: Optional function to add hyperlinks
    
    Returns:
        True if video was added, False otherwise
    """
    videos = get_video_for_attraction(attraction_name, city_name, max_videos=1)
    
    if not videos:
        return False
    
    video = videos[0]
    hyperlink_fn = add_hyperlink_func or add_hyperlink
    
    # Create compact video link
    video_para = doc.add_paragraph()
    video_para.paragraph_format.left_indent = Inches(0.5)
    
    yt_run = video_para.add_run('🎬 ')
    yt_run.font.size = Pt(9)
    
    watch_url = video.get('watch_url', '')
    hyperlink_fn(video_para, watch_url, 'Watch video preview')
    
    return True


# ============================================================================
# TESTING
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("YouTube Helper - Testing")
    print("=" * 60)
    
    # Test loading
    db = _load_youtube_db()
    print(f"\nLoaded {len(db)} destinations")
    
    # Test city lookup
    test_cities = ["Seville", "Granada", "Ronda", "Córdoba", "Málaga"]
    
    print("\n--- City Video Lookup ---")
    for city in test_cities:
        videos = get_video_for_city(city)
        if videos:
            print(f"✅ {city}: {videos[0]['title'][:50]}...")
        else:
            print(f"❌ {city}: No video found")
    
    # Test attraction lookup
    test_attractions = ["Alhambra", "Mezquita", "Caminito del Rey", "Plaza de España"]
    
    print("\n--- Attraction Video Lookup ---")
    for attr in test_attractions:
        videos = get_video_for_attraction(attr)
        if videos:
            print(f"✅ {attr}: {videos[0]['title'][:50]}...")
        else:
            print(f"❌ {attr}: No video found")
    
    print("\n" + "=" * 60)

# ============================================================================
# STREAMLIT UI FUNCTIONS (Moved from youtube_ui.py)
# ============================================================================

import streamlit as st
# Other required imports (json, os, unicodedata) are already at the top of youtube_helper.py

def convert_to_embed_url(youtube_url: str) -> str:
    """
    Converts a standard YouTube watch URL to an embed URL for Streamlit's st.video().
    This is necessary for videos to embed and play correctly in the UI.
    """
    if not youtube_url:
        return ""
    
    # 1. Handle standard 'watch' link
    if 'watch?v=' in youtube_url:
        # Replace 'watch?v=' with 'embed/' and strip any extra parameters like &t=
        return youtube_url.replace('watch?v=', 'embed/').split('&')[0]
    
    # 2. Handle short 'youtu.be' link
    elif 'youtu.be/' in youtube_url:
        # Extract the video ID after the last slash
        video_id = youtube_url.split('/')[-1].split('?')[0]
        return f"https://www.youtube.com/embed/{video_id}"
    
    # 3. Fallback
    return youtube_url

def display_video_expander(city_name: str, expanded: bool = False):
    """
    Display YouTube video inside an expander (collapsible) in the Streamlit UI.
    
    Args:
        city_name: City to show video for
        expanded: Whether expander is open by default
    """
    # Use the existing data-fetching function from youtube_helper.py
    videos = get_video_for_city(city_name, max_videos=1)
    
    if not videos:
        return False
    
    video = videos[0]
    watch_url = video.get('watch_url', '')
    title = video.get('title', 'Video Preview')
    channel = video.get('channel', 'Unknown Channel')
    
    if not watch_url:
        return False
    
    # Convert to embed URL before passing to st.video
    embed_url = convert_to_embed_url(watch_url)
    
    with st.expander(f"🎬 Watch: {title[:50]}...", expanded=expanded):
        st.video(embed_url)
        st.caption(f"📺 Channel: {channel}")
        st.link_button("▶️ Open in YouTube", watch_url, use_container_width=True)
    
    return True